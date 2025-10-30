"""
Generate DOCX Template for Akt o proceni rizika
Uses python-docx to create template with Mustache placeholders

Install: pip install python-docx
Run: python scripts/generate-docx-template.py
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_background(cell, color):
    """Set cell background color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

def create_template():
    # Create document
    doc = Document()

    # Set default font to Noto Sans
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Noto Sans'
    font.size = Pt(11)

    # =========================================================================
    # 1. NASLOVNA STRANA (Cover Page)
    # =========================================================================

    # Title - centered
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('РЕПУБЛИКА СРБИЈА')
    run.font.size = Pt(14)
    run.font.bold = True

    doc.add_paragraph()  # Blank line

    # Main title
    main_title = doc.add_paragraph()
    main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = main_title.add_run('АКТ О ПРОЦЕНИ РИЗИКА')
    run.font.size = Pt(18)
    run.font.bold = True

    doc.add_paragraph()

    # Company info - centered
    company_info = doc.add_paragraph()
    company_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = company_info.add_run('{{company.name}}\n')
    run.font.size = Pt(14)
    run.font.bold = True
    company_info.add_run('{{company.address}}, {{company.city}}\n')
    company_info.add_run('ПИБ: {{company.pib}}')

    doc.add_paragraph()
    doc.add_paragraph()

    # Date info - centered
    date_info = doc.add_paragraph()
    date_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_info.add_run('Датум израде: {{metadata.generatedDate}}\n')
    date_info.add_run('Важи до: {{metadata.validUntil}}')

    # Page break
    doc.add_page_break()

    # =========================================================================
    # 2. УВОД (Introduction)
    # =========================================================================

    heading = doc.add_heading('1. УВОД', level=1)
    heading.runs[0].font.name = 'Noto Sans'

    intro = doc.add_paragraph(
        'Овај Акт о процени ризика израђен је у складу са:'
    )

    doc.add_paragraph(
        '• Законом о безбедности и здрављу на раду („Сл. гласник РС", бр. 101/2005, 91/2015 и 113/2017)',
        style='List Bullet'
    )
    doc.add_paragraph(
        '• Правилником о начину и поступку процене ризика на радном месту и у радној околини („Сл. гласник РС", бр. 5/2018)',
        style='List Bullet'
    )

    doc.add_paragraph()

    doc.add_paragraph(
        f'Процена ризика обављена је за предузеће {{{{company.name}}}} '
        f'коришћењем Е×П×Ф методологије (Изложеност × Вероватноћа × Учесталост). '
        f'Циљ процене ризика је идентификација опасности и штетности на радним местима, '
        f'процена нивоа ризика и утврђивање мера за њихово отклањање или смањење.'
    )

    doc.add_paragraph()

    doc.add_paragraph(
        'Формула за израчунавање нивоа ризика:'
    )

    formula = doc.add_paragraph()
    formula.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = formula.add_run('Р = Е × П × Ф')
    run.font.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph()

    doc.add_paragraph(
        'Категорије ризика:\n'
        '• Низак ниво ризика (Р ≤ 36)\n'
        '• Средњи ниво ризика (36 < Р ≤ 70)\n'
        '• Висок ниво ризика (Р > 70)'
    )

    doc.add_page_break()

    # =========================================================================
    # 3. ПОДАЦИ О ПОСЛОДАВЦУ (Company Data)
    # =========================================================================

    heading = doc.add_heading('2. ПОДАЦИ О ПОСЛОДАВЦУ', level=1)
    heading.runs[0].font.name = 'Noto Sans'

    doc.add_paragraph(f'Пун назив: {{{{company.name}}}}')
    doc.add_paragraph(f'ПИБ: {{{{company.pib}}}}')
    doc.add_paragraph(f'Матични број: {{{{company.maticniBroj}}}}')
    doc.add_paragraph(f'Адреса седишта: {{{{company.address}}}}, {{{{company.city}}}}')
    doc.add_paragraph(f'Шифра делатности: {{{{company.activityCode}}}}')
    doc.add_paragraph(f'Опис делатности: {{{{company.activityDescription}}}}')
    doc.add_paragraph(f'Директор: {{{{company.director}}}}')
    doc.add_paragraph(f'Лице одговорно за БЗР: {{{{company.bzrResponsiblePerson}}}}')
    doc.add_paragraph(f'Укупан број запослених: {{{{company.employeeCount}}}}')

    doc.add_page_break()

    # =========================================================================
    # 4. ОРГАНИЗАЦИОНА СТРУКТУРА (Organizational Structure)
    # =========================================================================

    heading = doc.add_heading('3. ОРГАНИЗАЦИОНА СТРУКТУРА', level=1)
    heading.runs[0].font.name = 'Noto Sans'

    doc.add_paragraph(
        f'Предузеће {{{{company.name}}}} запошљава укупно {{{{summary.totalPositions}}}} '
        f'радника на {{{{summary.totalPositions}}}} радних места.'
    )

    doc.add_page_break()

    # =========================================================================
    # 5. СИСТЕМАТИЗАЦИЈА РАДНИХ МЕСТА (Position Systematization)
    # =========================================================================

    heading = doc.add_heading('4. СИСТЕМАТИЗАЦИЈА РАДНИХ МЕСТА', level=1)
    heading.runs[0].font.name = 'Noto Sans'

    # Position loop start
    doc.add_paragraph('{{#positions}}')

    doc.add_paragraph()
    pos_heading = doc.add_paragraph()
    run = pos_heading.add_run('4.{{@index}}. {{positionName}}')
    run.font.bold = True
    run.font.size = Pt(12)

    doc.add_paragraph('   • Организациона јединица: {{department}}')
    doc.add_paragraph('   • Број извршилаца: {{totalCount}} (М: {{maleCount}}, Ж: {{femaleCount}})')
    doc.add_paragraph('   • Захтевана школска спрема: {{requiredEducation}}')
    doc.add_paragraph('   • Потребно искуство: {{requiredExperience}}')
    doc.add_paragraph('   • Опис послова: {{workDescription}}')

    doc.add_paragraph('{{/positions}}')

    doc.add_page_break()

    # =========================================================================
    # 6. ПРОЦЕНА РИЗИКА ПО РАДНИМ МЕСТИМА (Risk Assessment)
    # =========================================================================

    heading = doc.add_heading('5. ПРОЦЕНА РИЗИКА ПО РАДНИМ МЕСТИМА', level=1)
    heading.runs[0].font.name = 'Noto Sans'

    # Position loop
    doc.add_paragraph('{{#positions}}')

    doc.add_paragraph()
    pos_heading = doc.add_paragraph()
    run = pos_heading.add_run('5.{{@index}}. Радно место: {{positionName}}')
    run.font.bold = True
    run.font.size = Pt(12)

    doc.add_paragraph()

    # Create risk assessment table
    table = doc.add_table(rows=2, cols=11)
    table.style = 'Table Grid'

    # Header row
    headers = ['Опасност/Штетност', 'Еи', 'Пи', 'Фи', 'Ри', 'Мере превенције', 'Е', 'П', 'Ф', 'Р', 'Ниво']
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        cell = header_cells[i]
        cell.text = header
        # Bold and centered
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(9)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Gray background
        set_cell_background(cell, 'D9D9D9')

    # Data row with Mustache placeholders
    doc.add_paragraph('{{#risks}}')

    data_row = table.rows[1]
    placeholders = [
        '{{hazardName}}',
        '{{ei}}',
        '{{pi}}',
        '{{fi}}',
        '{{ri}}',
        '{{correctiveMeasures}}',
        '{{e}}',
        '{{p}}',
        '{{f}}',
        '{{r}}',
        '{{riskLevel}}'
    ]

    for i, placeholder in enumerate(placeholders):
        cell = data_row.cells[i]
        cell.text = placeholder
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)
            if i > 0 and i != 5:  # Center numbers, not text
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('{{/risks}}')

    doc.add_paragraph()
    note = doc.add_paragraph()
    run = note.add_run('Напомена: ')
    run.font.bold = True
    note.add_run('Ниво ризика - Низак (Р ≤ 36), Средњи (36 < Р ≤ 70), Висок (Р > 70)')

    doc.add_paragraph('{{/positions}}')

    doc.add_page_break()

    # =========================================================================
    # 7. ЗБОРНИ ПРИКАЗ (Summary)
    # =========================================================================

    heading = doc.add_heading('6. ЗБОРНИ ПРИКАЗ ПРОЦЕНЕ', level=1)
    heading.runs[0].font.name = 'Noto Sans'

    doc.add_paragraph(f'Укупно радних места: {{{{summary.totalPositions}}}}')
    doc.add_paragraph(f'Укупно процењених опасности: {{{{summary.totalRisks}}}}')

    doc.add_paragraph()

    summary_heading = doc.add_paragraph()
    run = summary_heading.add_run('Расподела ризика:')
    run.font.bold = True

    doc.add_paragraph(f'• Низак ниво (Р ≤ 36): {{{{summary.lowRiskCount}}}}')
    doc.add_paragraph(f'• Средњи ниво (36 < Р ≤ 70): {{{{summary.mediumRiskCount}}}}')
    doc.add_paragraph(f'• Висок ниво (Р > 70): {{{{summary.highRiskCount}}}}')

    doc.add_paragraph()

    # Warning for high risks
    doc.add_paragraph('{{#summary.highRiskPositions}}')
    warning = doc.add_paragraph()
    run = warning.add_run('⚠️ Радна места са високим ризиком: {{summary.highRiskPositions}}')
    run.font.color.rgb = RGBColor(255, 0, 0)
    run.font.bold = True
    doc.add_paragraph('{{/summary.highRiskPositions}}')

    doc.add_page_break()

    # =========================================================================
    # 8. ПРИЛОЗИ (Appendices)
    # =========================================================================

    heading = doc.add_heading('7. ПРИЛОЗИ', level=1)
    heading.runs[0].font.name = 'Noto Sans'

    doc.add_paragraph(
        'Прилози који прате овај Акт о процени ризика:\n'
        '• Списак потребних средстава и опреме за личну заштиту\n'
        '• План обуке запослених за безбедан рад\n'
        '• Списак лекарских прегледа'
    )

    doc.add_page_break()

    # =========================================================================
    # 9. ВЕРИФИКАЦИЈА (Verification)
    # =========================================================================

    heading = doc.add_heading('8. ВЕРИФИКАЦИЈА И ПОТПИСИ', level=1)
    heading.runs[0].font.name = 'Noto Sans'

    doc.add_paragraph(
        f'Овај Акт о процени ризика верификован је и ступа на снагу даном {{{{metadata.generatedDate}}}}.'
    )

    doc.add_paragraph(
        f'Важи {{{{metadata.validityPeriod}}}} од дана доношења (у складу са чланом 32 Правилника).'
    )

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    # Signature block - Director
    director_line = doc.add_paragraph('_' * 30)
    director_line.alignment = WD_ALIGN_PARAGRAPH.LEFT

    director_name = doc.add_paragraph('{{company.director}}')
    director_name.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = director_name.runs[0]
    run.font.bold = True

    director_title = doc.add_paragraph('Директор')
    director_title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph()
    doc.add_paragraph()

    # Signature block - BZR Responsible
    bzr_line = doc.add_paragraph('_' * 30)
    bzr_line.alignment = WD_ALIGN_PARAGRAPH.LEFT

    bzr_name = doc.add_paragraph('{{company.bzrResponsiblePerson}}')
    bzr_name.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = bzr_name.runs[0]
    run.font.bold = True

    bzr_title = doc.add_paragraph('Лице одговорно за безбедност и здравље на раду')
    bzr_title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Footer
    doc.add_paragraph()
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run('Документ генерисан аутоматски системом BZR Portal')
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(128, 128, 128)

    # Save
    output_path = 'templates/akt-procena-rizika-template.docx'
    doc.save(output_path)
    print(f'[OK] Template created: {output_path}')
    print(f'[INFO] Total pages: ~{len(doc.sections)} sections')
    print(f'[INFO] Ready for docx-templates library')
    print()
    print('Next: Test with npm test -- document.service.test.ts')

if __name__ == '__main__':
    try:
        create_template()
    except Exception as e:
        print(f'[ERROR] Error: {e}')
        print()
        print('Make sure python-docx is installed:')
        print('  pip install python-docx')
