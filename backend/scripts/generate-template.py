#!/usr/bin/env python3
"""
Generate BZR Portal DOCX Template with Mustache placeholders
Based on TEMPLATE_SPECIFICATION.md

Requires: python-docx
Install: pip install python-docx
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_header_text(paragraph, text, font_size=14, bold=True, align='left'):
    """Add formatted header text"""
    run = paragraph.add_run(text)
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.name = 'Arial'

    if align == 'center':
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right':
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

def add_paragraph_with_label(doc, label, placeholder, bold_label=True):
    """Add paragraph with bold label and placeholder"""
    p = doc.add_paragraph()
    if bold_label:
        run = p.add_run(f"{label}: ")
        run.font.bold = True
        run.font.size = Pt(12)
        run.font.name = 'Arial'

    run = p.add_run(placeholder)
    run.font.size = Pt(12)
    run.font.name = 'Arial'

def create_risk_table(doc):
    """Create risk assessment table with Mustache loop"""
    # Add table explanation
    p = doc.add_paragraph()
    run = p.add_run("Табела процене ризика:")
    run.font.bold = True
    run.font.size = Pt(12)
    doc.add_paragraph()

    # Table header
    table = doc.add_table(rows=1, cols=13)
    table.style = 'Light Grid Accent 1'

    headers = ['Р.Бр.', 'Опасност', 'Категорија', 'E', 'П', 'Ф', 'Ri',
               'Корективне мере', 'E', 'П', 'Ф', 'R', 'Ниво ризика']

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        # Bold header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)

    # Add Mustache loop row
    row = table.add_row()
    placeholders = [
        '{{rowNumber}}',
        '{{hazardName}}',
        '{{hazardCategory}}',
        '{{ei}}',
        '{{pi}}',
        '{{fi}}',
        '{{ri}}',
        '{{correctiveMeasures}}',
        '{{e}}',
        '{{p}}',
        '{{f}}',
        '{{r}}',
        '{{riskLevel}}'
    ]

    for i, placeholder in enumerate(placeholders):
        row.cells[i].text = placeholder

def generate_template():
    """Generate complete DOCX template"""
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)

    # =============================================================================
    # 1. NASLOV (Title Page)
    # =============================================================================
    title = doc.add_paragraph()
    add_header_text(title, "АКТ О ПРОЦЕНИ РИЗИКА", font_size=18, bold=True, align='center')
    doc.add_paragraph()

    add_paragraph_with_label(doc, "Предузеће", "{{company.name}}")
    add_paragraph_with_label(doc, "ПИБ", "{{company.pib}}")
    add_paragraph_with_label(doc, "Адреса", "{{company.address}}")
    doc.add_paragraph()
    add_paragraph_with_label(doc, "Датум израде", "{{metadata.generatedDate}}")
    add_paragraph_with_label(doc, "Важност акта", "{{metadata.validityPeriod}}")

    doc.add_page_break()

    # =============================================================================
    # 2. ОПШТИ ПОДАЦИ О ПОСЛОДАВЦУ (Company Information)
    # =============================================================================
    section1 = doc.add_paragraph()
    add_header_text(section1, "I. ОПШТИ ПОДАЦИ О ПОСЛОДАВЦУ", font_size=14, bold=True)
    doc.add_paragraph()

    add_paragraph_with_label(doc, "Назив предузећа", "{{company.name}}")
    add_paragraph_with_label(doc, "ПИБ", "{{company.pib}}")
    add_paragraph_with_label(doc, "Адреса седишта", "{{company.address}}")
    add_paragraph_with_label(doc, "Град", "{{company.city}}")
    add_paragraph_with_label(doc, "Шифра делатности", "{{company.activityCode}}")
    add_paragraph_with_label(doc, "Опис делатности", "{{company.activityDescription}}")
    add_paragraph_with_label(doc, "Број запослених", "{{company.employeeCount}}")
    doc.add_paragraph()
    add_paragraph_with_label(doc, "Директор", "{{company.director}}")
    add_paragraph_with_label(doc, "Лице задужено за БЗР", "{{company.bzrResponsiblePerson}}")

    doc.add_page_break()

    # =============================================================================
    # 3. СИСТЕМАТИЗАЦИЈА РАДНИХ МЕСТА (Position Systematization)
    # =============================================================================
    section2 = doc.add_paragraph()
    add_header_text(section2, "II. СИСТЕМАТИЗАЦИЈА РАДНИХ МЕСТА", font_size=14, bold=True)
    doc.add_paragraph()

    # Mustache loop for positions
    doc.add_paragraph("{{#positions}}")
    doc.add_paragraph()

    add_header_text(doc.add_paragraph(), "Радно место: {{positionName}}", font_size=12, bold=True)
    add_paragraph_with_label(doc, "Организациона јединица", "{{department}}")

    p = doc.add_paragraph()
    run = p.add_run("Број запослених: ")
    run.font.bold = True
    run.font.size = Pt(12)
    p.add_run(f"Укупно: {{{{totalCount}}}}, Мушкарци: {{{{maleCount}}}}, Жене: {{{{femaleCount}}}}")

    add_paragraph_with_label(doc, "Потребна стручна спрема", "{{requiredEducation}}")
    add_paragraph_with_label(doc, "Радно искуство", "{{requiredExperience}}")

    doc.add_paragraph()
    doc.add_paragraph("{{/positions}}")

    doc.add_page_break()

    # =============================================================================
    # 4. ПРОЦЕНА РИЗИКА (Risk Assessment)
    # =============================================================================
    section3 = doc.add_paragraph()
    add_header_text(section3, "III. ПРОЦЕНА РИЗИКА ПО РАДНИМ МЕСТИМА", font_size=14, bold=True)
    doc.add_paragraph()

    # Add loop for positions with risks
    doc.add_paragraph("{{#positions}}")
    add_header_text(doc.add_paragraph(), "Радно место: {{positionName}}", font_size=12, bold=True)
    doc.add_paragraph()

    # Add paragraph explaining E×P×F methodology
    methodology = doc.add_paragraph()
    run = methodology.add_run("Методологија процене: ")
    run.font.bold = True
    methodology.add_run("E (Последице) × П (Вероватноћа) × Ф (Учесталост)")
    doc.add_paragraph()

    # Add nested loop for risks
    doc.add_paragraph("{{#risks}}")

    # Risk entry box
    risk_box = doc.add_paragraph()
    run = risk_box.add_run("━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━")

    add_paragraph_with_label(doc, "Опасност", "{{hazardName}} (Код: {{hazardCode}})")
    add_paragraph_with_label(doc, "Категорија", "{{hazardCategory}}")
    doc.add_paragraph()

    # Initial risk
    initial = doc.add_paragraph()
    run = initial.add_run("ПОЧЕТНА ПРОЦЕНА (пре корективних мера):")
    run.font.bold = True
    run.font.size = Pt(11)

    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f"E (Последице): {{{{ei}}}}, П (Вероватноћа): {{{{pi}}}}, Ф (Учесталост): {{{{fi}}}}")

    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run("Ri (Почетни ризик): ")
    run.font.bold = True
    p.add_run("{{ri}} {{initialRiskLevel}}")

    doc.add_paragraph()

    # Corrective measures
    measures = doc.add_paragraph()
    run = measures.add_run("КОРЕКТИВНЕ МЕРЕ:")
    run.font.bold = True
    run.font.size = Pt(11)

    doc.add_paragraph("{{correctiveMeasures}}")
    doc.add_paragraph()

    # Residual risk
    residual = doc.add_paragraph()
    run = residual.add_run("РЕЗИДУАЛНА ПРОЦЕНА (након корективних мера):")
    run.font.bold = True
    run.font.size = Pt(11)

    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f"E (Последице): {{{{e}}}}, П (Вероватноћа): {{{{p}}}}, Ф (Учесталост): {{{{f}}}}")

    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run("R (Резидуални ризик): ")
    run.font.bold = True
    p.add_run("{{r}} {{riskLevel}}")

    # Check if high risk
    doc.add_paragraph()
    doc.add_paragraph("{{#isHighRisk}}")
    warning = doc.add_paragraph()
    run = warning.add_run("⚠️ ВИСОК РИЗИК - Потребна хитна интервенција!")
    run.font.bold = True
    run.font.color.rgb = RGBColor(220, 53, 69)  # Red
    doc.add_paragraph("{{/isHighRisk}}")

    risk_box2 = doc.add_paragraph()
    run = risk_box2.add_run("━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━")
    doc.add_paragraph()

    doc.add_paragraph("{{/risks}}")
    doc.add_paragraph()
    doc.add_paragraph("{{/positions}}")

    doc.add_page_break()

    # =============================================================================
    # 5. ЗБИРНИ ПРЕГЛЕД (Summary)
    # =============================================================================
    section4 = doc.add_paragraph()
    add_header_text(section4, "IV. ЗБИРНИ ПРЕГЛЕД", font_size=14, bold=True)
    doc.add_paragraph()

    summary_intro = doc.add_paragraph()
    summary_intro.add_run("На основу спроведене процене ризика утврђено је следеће:")
    doc.add_paragraph()

    add_paragraph_with_label(doc, "Укупан број радних места", "{{summary.totalPositions}}")
    add_paragraph_with_label(doc, "Укупан број идентификованих ризика", "{{summary.totalRisks}}")
    add_paragraph_with_label(doc, "Радна места са високим ризиком", "{{summary.highRiskPositions}}")
    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("Распоред ризика по нивоима:")
    run.font.bold = True

    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f"Низак ризик (≤36): {{{{summary.lowRiskCount}}}}")

    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f"Средњи ризик (37-70): {{{{summary.mediumRiskCount}}}}")

    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f"Висок ризик (>70): {{{{summary.highRiskCount}}}}")

    doc.add_paragraph()

    conclusion = doc.add_paragraph()
    conclusion.add_run(
        "Све идентификоване опасности су проценене према методологији E×П×Ф "
        "(Последице × Вероватноћа × Учесталост) у складу са Законом о безбедности "
        "и здрављу на раду (Сл. гласник РС, бр. 101/2005...91/2015)."
    )

    doc.add_page_break()

    # =============================================================================
    # 6. ПОТПИСНИЦИ (Signatures)
    # =============================================================================
    section5 = doc.add_paragraph()
    add_header_text(section5, "V. ПОТПИСНИЦИ", font_size=14, bold=True)
    doc.add_paragraph()
    doc.add_paragraph()

    # Signature lines
    sig1 = doc.add_paragraph()
    sig1.add_run("Директор: {{company.director}}")
    doc.add_paragraph()
    sig1_line = doc.add_paragraph()
    sig1_line.add_run("_" * 40)
    sig1_line.add_run("  Потпис")

    doc.add_paragraph()
    doc.add_paragraph()

    sig2 = doc.add_paragraph()
    sig2.add_run("Лице задужено за БЗР: {{company.bzrResponsiblePerson}}")
    doc.add_paragraph()
    sig2_line = doc.add_paragraph()
    sig2_line.add_run("_" * 40)
    sig2_line.add_run("  Потпис")

    doc.add_paragraph()
    doc.add_paragraph()

    # Date line
    date_line = doc.add_paragraph()
    date_line.add_run("Датум: {{metadata.generatedDate}}")
    doc.add_paragraph()
    doc.add_paragraph()
    date_sig = doc.add_paragraph()
    date_sig.add_run("_" * 40)

    # =============================================================================
    # Save document
    # =============================================================================
    output_path = "../templates/akt-procena-rizika-template.docx"
    doc.save(output_path)
    print(f"✅ Template generated successfully: {output_path}")
    print(f"   File contains {len(doc.paragraphs)} paragraphs")
    print(f"   Ready for docx-templates library with Mustache syntax")

if __name__ == "__main__":
    generate_template()
