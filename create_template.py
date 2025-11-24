#!/usr/bin/env python3
"""
Kreiranje DOCX template fajla sa Mustache placeholders
Za BZR Portal - Akt o proceni rizika
Prema FR-034 do FR-042 zahtevima
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Kreiranje novog Word dokumenta
doc = Document()

# Podešavanje margina (2.5cm)
sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.98)
    section.bottom_margin = Inches(0.98)
    section.left_margin = Inches(0.98)
    section.right_margin = Inches(0.98)

# ==========================================
# NASLOVNA STRANA (FR-034)
# ==========================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('АКТ О ПРОЦЕНИ РИЗИКА')
run.bold = True
run.font.size = Pt(18)
run.font.name = 'Times New Roman'

doc.add_paragraph()

company_name = doc.add_paragraph()
company_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = company_name.add_run('{{company.name}}')
run.bold = True
run.font.size = Pt(14)
run.font.name = 'Times New Roman'

doc.add_paragraph()
address_p = doc.add_paragraph()
address_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = address_p.add_run('{{company.address}}')
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

doc.add_page_break()

# ==========================================
# 1. UVOD (FR-035)
# ==========================================
heading = doc.add_heading('1. УВОД', level=1)
heading.runs[0].font.name = 'Times New Roman'

intro = doc.add_paragraph(
    'Овај Акт о процени ризика сачињен је у складу са Законом о безбедности и здрављу на раду '
    '(„Службени гласник РС", бр. 101/2005, 91/2015 и 113/2017) и Правилником о превентивним мерама '
    'за безбедан и здрав рад при излагању ризицима услед појаве штетности у радној околини '
    '(„Службени гласник РС", бр. 5/2018).'
)
intro.runs[0].font.name = 'Times New Roman'
intro.runs[0].font.size = Pt(11)

purpose = doc.add_paragraph(
    'Циљ процене ризика је идентификација опасности, процена ризика по безбедност и здравље запослених, '
    'утврђивање мера за отклањање ризика или њихово свођење на прихватљив ниво.'
)
purpose.runs[0].font.name = 'Times New Roman'
purpose.runs[0].font.size = Pt(11)

# ==========================================
# 2. PODACI O POSLODAVCU (FR-036)
# ==========================================
heading2 = doc.add_heading('2. ПОДАЦИ О ПОСЛОДАВЦУ', level=1)
heading2.runs[0].font.name = 'Times New Roman'

employer_table = doc.add_table(rows=7, cols=2)
employer_table.style = 'Table Grid'

cells_data = [
    ('Назив:', '{{company.name}}'),
    ('Адреса:', '{{company.address}}'),
    ('ПИБ:', '{{company.pib}}'),
    ('Матични број:', '{{company.registration_number}}'),
    ('Шифра делатности:', '{{company.activity_code}}'),
    ('Директор:', '{{company.director}}'),
    ('Одговорно лице за БЗР:', '{{company.bzr_responsible_person}}'),
]

for i, (label, value) in enumerate(cells_data):
    row = employer_table.rows[i]
    row.cells[0].text = label
    row.cells[1].text = value
    for para in row.cells[0].paragraphs:
        for run in para.runs:
            run.bold = True
            run.font.name = 'Times New Roman'
    for para in row.cells[1].paragraphs:
        for run in para.runs:
            run.font.name = 'Times New Roman'

doc.add_page_break()

# ==========================================
# 3. SISTEMATIZACIJA RADNIH MESTA (FR-038)
# ==========================================
heading4 = doc.add_heading('3. СИСТЕМАТИЗАЦИЈА РАДНИХ МЕСТА', level=1)
heading4.runs[0].font.name = 'Times New Roman'

sys_intro = doc.add_paragraph('Табеларни преглед свих радних места:')
sys_intro.runs[0].font.name = 'Times New Roman'
doc.add_paragraph()

# Mustache loop za radna mesta
loop_start = doc.add_paragraph('{{#positions}}')

position_heading = doc.add_heading('Радно место: {{position_name}}', level=2)
position_heading.runs[0].font.name = 'Times New Roman'

position_table = doc.add_table(rows=7, cols=2)
position_table.style = 'Table Grid'

position_data = [
    ('Шифра:', '{{position_code}}'),
    ('Организациона јединица:', '{{department}}'),
    ('Потребна стручна спрема:', '{{required_education}}'),
    ('Број извршилаца (М/Ж):', '{{employees_male}}/{{employees_female}}'),
    ('Радни сати (дневно):', '{{work_hours_daily}}'),
    ('Опис послова:', '{{job_description}}'),
    ('Радни простор:', '{{workspace}}'),
]

for i, (label, value) in enumerate(position_data):
    row = position_table.rows[i]
    row.cells[0].text = label
    row.cells[1].text = value
    for para in row.cells[0].paragraphs:
        for run in para.runs:
            run.bold = True
            run.font.name = 'Times New Roman'
    for para in row.cells[1].paragraphs:
        for run in para.runs:
            run.font.name = 'Times New Roman'

doc.add_paragraph()
loop_end = doc.add_paragraph('{{/positions}}')

doc.add_page_break()

# ==========================================
# 4. PROCENA RIZIKA (FR-039)
# ==========================================
heading5 = doc.add_heading('4. ПРОЦЕНА РИЗИКА ПО РАДНИМ МЕСТИМА', level=1)
heading5.runs[0].font.name = 'Times New Roman'

methodology = doc.add_paragraph(
    'Процена ризика спроведена је применом методологије Е × П × Ф:\n\n'
    '• Е (Ефекат) - озбиљност последице (1-6)\n'
    '• П (Вероватноћа) - вероватноћа настанка штетног догађаја (1-6)\n'
    '• Ф (Фреквенција) - учесталост излагања опасности (1-6)\n'
    '• Ri (Иницијални ризик) = Е × П × Ф (пре примене мера)\n'
    '• R (Резидуални ризик) = Е × П × Ф (након примене мера)'
)
methodology.runs[0].font.name = 'Times New Roman'
methodology.runs[0].font.size = Pt(10)

doc.add_paragraph()

# Loop pozicija i rizika
positions_loop = doc.add_paragraph('{{#positions}}')

pos_heading = doc.add_heading('Радно место: {{position_name}}', level=2)
pos_heading.runs[0].font.name = 'Times New Roman'

risks_loop = doc.add_paragraph('{{#risks}}')

risk_table = doc.add_table(rows=16, cols=2)
risk_table.style = 'Table Grid'

risk_data = [
    ('Опасност:', '{{hazard.hazard_code}} - {{hazard.hazard_name_sr}}'),
    ('', ''),
    ('ИНИЦИЈАЛНИ РИЗИК:', ''),
    ('  Е (иницијални):', '{{initial_e}}'),
    ('  П (иницијална):', '{{initial_p}}'),
    ('  Ф (иницијална):', '{{initial_f}}'),
    ('  Ri = Е × П × Ф:', '{{initial_ri}}'),
    ('', ''),
    ('Корективне мере:', '{{corrective_measures}}'),
    ('', ''),
    ('РЕЗИДУАЛНИ РИЗИК:', ''),
    ('  Е (резидуални):', '{{residual_e}}'),
    ('  П (резидуална):', '{{residual_p}}'),
    ('  Ф (резидуална):', '{{residual_f}}'),
    ('  R = Е × П × Ф:', '{{residual_r}}'),
    ('Одговорно лице:', '{{responsible_person}}'),
]

for i, (label, value) in enumerate(risk_data):
    row = risk_table.rows[i]
    row.cells[0].text = label
    row.cells[1].text = value
    if label and not label.startswith('  '):
        for para in row.cells[0].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.name = 'Times New Roman'
    for para in row.cells[1].paragraphs:
        for run in para.runs:
            run.font.name = 'Times New Roman'

doc.add_paragraph()
risks_end = doc.add_paragraph('{{/risks}}')
positions_end = doc.add_paragraph('{{/positions}}')

doc.add_page_break()

# ==========================================
# 5. ZBIRNI PRIKAZ (FR-040)
# ==========================================
heading6 = doc.add_heading('5. ЗБИРНИ ПРИКАЗ ПРОЦЕНЕ РИЗИКА', level=1)
heading6.runs[0].font.name = 'Times New Roman'

summary_intro = doc.add_paragraph('Табеларни преглед свих процењених ризика:')
summary_intro.runs[0].font.name = 'Times New Roman'

summary_table = doc.add_table(rows=1, cols=5)
summary_table.style = 'Table Grid'

headers = ['Радно место', 'Опасност', 'Ri', 'R', 'Ниво ризика']
for i, header in enumerate(headers):
    cell = summary_table.rows[0].cells[i]
    cell.text = header
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)

doc.add_paragraph()
summary_loop = doc.add_paragraph('{{#summary_risks}}')
summary_row = doc.add_paragraph('{{position_name}} | {{hazard_name}} | {{initial_ri}} | {{residual_r}} | {{risk_level}}')
summary_row.runs[0].font.size = Pt(10)
summary_row.runs[0].font.name = 'Times New Roman'
summary_end = doc.add_paragraph('{{/summary_risks}}')

doc.add_page_break()

# ==========================================
# 6. LZO (FR-041a)
# ==========================================
heading7 = doc.add_heading('6. ЛИЧНА ЗАШТИТНА ОПРЕМА', level=1)
heading7.runs[0].font.name = 'Times New Roman'

ppe_intro = doc.add_paragraph('Утврђене потребе за личном заштитном опремом:')
ppe_intro.runs[0].font.name = 'Times New Roman'

doc.add_paragraph('{{#positions}}')
doc.add_paragraph('{{#has_ppe}}')

ppe_pos_heading = doc.add_heading('Радно место: {{position_name}}', level=3)
ppe_pos_heading.runs[0].font.name = 'Times New Roman'

doc.add_paragraph('{{#ppe_items}}')
ppe_item = doc.add_paragraph('• {{ppe_type}} ({{ppe_standard}}) - {{quantity}} ком.')
ppe_item.runs[0].font.name = 'Times New Roman'
doc.add_paragraph('{{/ppe_items}}')

doc.add_paragraph('{{/has_ppe}}')
doc.add_paragraph('{{/positions}}')

doc.add_page_break()

# ==========================================
# 7. OBUKA (FR-041b)
# ==========================================
heading8 = doc.add_heading('7. ОБУКА ЗАПОСЛЕНИХ', level=1)
heading8.runs[0].font.name = 'Times New Roman'

training_intro = doc.add_paragraph('Утврђени захтеви за обуком запослених:')
training_intro.runs[0].font.name = 'Times New Roman'

doc.add_paragraph('{{#positions}}')
doc.add_paragraph('{{#has_training}}')

training_pos_heading = doc.add_heading('Радно место: {{position_name}}', level=3)
training_pos_heading.runs[0].font.name = 'Times New Roman'

doc.add_paragraph('{{#training_requirements}}')
training_item = doc.add_paragraph('• {{training_type}} - фреквенција: {{frequency_months}} месеци, трајање: {{duration_hours}}h')
training_item.runs[0].font.name = 'Times New Roman'
doc.add_paragraph('{{/training_requirements}}')

doc.add_paragraph('{{/has_training}}')
doc.add_paragraph('{{/positions}}')

doc.add_page_break()

# ==========================================
# 8. LEKARSKI PREGLEDI (FR-041c)
# ==========================================
heading9 = doc.add_heading('8. ЛЕКАРСКИ ПРЕГЛЕДИ', level=1)
heading9.runs[0].font.name = 'Times New Roman'

medical_intro = doc.add_paragraph('Утврђени захтеви за лекарским прегледима:')
medical_intro.runs[0].font.name = 'Times New Roman'

doc.add_paragraph('{{#positions}}')
doc.add_paragraph('{{#has_medical}}')

medical_pos_heading = doc.add_heading('Радно место: {{position_name}}', level=3)
medical_pos_heading.runs[0].font.name = 'Times New Roman'

doc.add_paragraph('{{#medical_exams}}')
medical_item = doc.add_paragraph('• {{exam_type}} - фреквенција: {{frequency_months}} месеци, обим: {{exam_scope}}')
medical_item.runs[0].font.name = 'Times New Roman'
doc.add_paragraph('{{/medical_exams}}')

doc.add_paragraph('{{/has_medical}}')
doc.add_paragraph('{{/positions}}')

doc.add_page_break()

# ==========================================
# 9. ZAKLJUCAK (FR-042a)
# ==========================================
heading10 = doc.add_heading('9. ЗАКЉУЧАК', level=1)
heading10.runs[0].font.name = 'Times New Roman'

conclusion = doc.add_paragraph(
    'Овај Акт о процени ризика обухвата све идентификоване опасности на радним местима '
    'послодавца {{company.name}}. Спроведеном проценом ризика утврђене су мере за безбедан и '
    'здрав рад, које је послодавац дужан да спроведе и континуирано прати њихову ефикасност.\n\n'
    'Ревизија Акта о процени ризика вршиће се периодично (минимум једном годишње), као и '
    'приликом значајних промена у технологији рада, организацији рада или након повреда на раду.'
)
conclusion.runs[0].font.name = 'Times New Roman'
conclusion.runs[0].font.size = Pt(11)

doc.add_paragraph()
doc.add_paragraph()

# ==========================================
# 10. POTPISI (FR-042)
# ==========================================
heading11 = doc.add_heading('10. ПОТПИСИ', level=1)
heading11.runs[0].font.name = 'Times New Roman'

signatures_text = doc.add_paragraph('Овај Акт о процени ризика сачинио је:')
signatures_text.runs[0].font.name = 'Times New Roman'

doc.add_paragraph()

signature_table = doc.add_table(rows=4, cols=2)
signature_table.style = 'Table Grid'

sig_data = [
    ('Састављач акта:', '{{company.bzr_responsible_person}}'),
    ('Потпис:', '_' * 40),
    ('Датум израде:', '{{generated_date}}'),
    ('Директор:', '{{company.director}}'),
]

for i, (label, value) in enumerate(sig_data):
    row = signature_table.rows[i]
    row.cells[0].text = label
    row.cells[1].text = value
    for para in row.cells[0].paragraphs:
        for run in para.runs:
            run.bold = True
            run.font.name = 'Times New Roman'
    for para in row.cells[1].paragraphs:
        for run in para.runs:
            run.font.name = 'Times New Roman'

doc.add_paragraph()
doc.add_paragraph()

final_note = doc.add_paragraph(
    'Напомена: Овај документ је генерисан аутоматски помоћу BZR Portal платформе.'
)
final_note.runs[0].font.name = 'Times New Roman'
final_note.runs[0].font.size = Pt(9)
final_note.runs[0].font.italic = True
final_note.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Sačuvaj dokument
output_path = 'backend/templates/Akt_Procena_Rizika_Template.docx'
doc.save(output_path)

print(f'Template successfully created: {output_path}')
print(f'Document contains:')
print(f'   - {len(doc.paragraphs)} paragraphs')
print(f'   - {len(doc.tables)} tables')
print(f'   - Mustache placeholders for dynamic data')
print(f'   - Serbian Cyrillic font (Times New Roman)')
print(f'   - All FR-034 to FR-042 sections included')
print(f'\nT084 COMPLETED!')
